# cv/build_template.py - one-time script that recreates the CV as a docx
# template with docxtpl placeholders in the Projects section.
# Run once: python cv/build_template.py  -> writes cv/cv_template.docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x40, 0x40, 0x40)


def para(doc, text="", size=10, bold=False, italic=False, color=None,
         align=None, space_after=2, space_before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def heading(doc, text):
    p = para(doc, text, size=11, bold=True, color=BLUE, space_before=8, space_after=3)
    # bottom border
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)

    # Header
    para(doc, "Ayush Das", size=22, bold=True, space_after=0)
    para(doc, "Email: das.ayush4890@gmail.com | Mobile: +91-9856912390", size=9, color=GRAY, space_after=0)
    para(doc, "Jalandhar, Punjab | Open to relocation & remote", size=9, color=GRAY, space_after=0)
    para(doc, "LinkedIn: linkedin.com/in/ayushdas4890 | GitHub: github.com/AyushDas4890", size=9, color=BLUE, space_after=4)

    # Skills
    heading(doc, "SKILLS")
    skills = [
        ("LLMs & RAG: ", "Transformers, Hugging Face, DeBERTa-v3, LangChain, LangGraph, sentence-transformers, Prompt Engineering"),
        ("Vector Search: ", "ChromaDB, FAISS, Embeddings, Hybrid Retrieval"),
        ("ML: ", "PyTorch, TensorFlow, Scikit-Learn, XGBoost, SHAP, Pandas, NumPy"),
        ("Deployment & MLOps: ", "FastAPI, REST APIs, Docker, LangSmith, LLMOps, Streamlit, Git, Hugging Face, SSE"),
        ("Languages: ", "Python, C++, SQL, JavaScript"),
        ("Soft Skills: ", "Cross-functional Collaboration, Technical Communication, Ownership, Fast Iteration, Adaptability"),
    ]
    for label, rest in skills:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label); r.bold = True; r.font.size = Pt(9.5)
        r2 = p.add_run(rest); r2.font.size = Pt(9.5)

    # Internship
    heading(doc, "INTERNSHIP")
    p = para(doc, "", space_after=1)
    r = p.add_run("Vanillakart"); r.bold = True; r.font.size = Pt(10)
    r = p.add_run("  |  Nov 2025 – Jan 2026"); r.font.size = Pt(9); r.font.color.rgb = GRAY
    para(doc, "WordPress Developer", size=9.5, italic=True, space_after=1)
    for b in [
        "Shipped and maintained production client websites on WordPress, owning responsive front-end development (HTML, CSS, JavaScript) end-to-end through deployment.",
        "Executed full content and asset migrations following SEO best practices, preserving all existing rankings with zero broken links post-launch.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(b); r.font.size = Pt(9.5)

    # PROJECTS — the templated section
    heading(doc, "PROJECTS")
    para(doc, "{%p for p in projects %}", size=9.5, space_after=1)
    p = para(doc, "", space_after=1)
    r = p.add_run("{{ p.title }}"); r.bold = True; r.font.size = Pt(10)
    p2 = para(doc, "", space_after=1)
    r = p2.add_run("Tech: "); r.bold = True; r.font.size = Pt(9)
    r = p2.add_run('{{ p.tech_stack | join(", ") }}'); r.font.size = Pt(9); r.italic = True
    para(doc, "{%p for b in p.bullets %}", size=9.5, space_after=1)
    pb = doc.add_paragraph(style="List Bullet")
    pb.paragraph_format.space_after = Pt(1)
    r = pb.add_run("{{ b }}"); r.font.size = Pt(9.5)
    para(doc, "{%p endfor %}", size=9.5, space_after=1)
    p3 = para(doc, "", space_after=3)
    r = p3.add_run("{{ p.repo_url }}"); r.font.size = Pt(8.5); r.font.color.rgb = BLUE
    para(doc, "{%p endfor %}", size=9.5, space_after=1)

    # Achievements
    heading(doc, "ACHIEVEMENTS")
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(
        "Contributed to OpenCV (open-source computer vision library) and authored an accepted "
        "documentation fix for LlamaIndex correcting a contradictory default model reference and a grammar error."
    ); r.font.size = Pt(9.5)

    # Certificates
    heading(doc, "CERTIFICATES")
    for c in [
        "IBM – Python for Data Science and AI | June 2025",
        "IBM – ChatGPT-4 Prompt Engineering | June 2025",
        "Udemy – Master Generative AI & Generative AI Tools | June 2025",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(c); r.font.size = Pt(9.5)

    # Education
    heading(doc, "EDUCATION")
    edu = [
        ("Lovely Professional University — Punjab, India",
         "Bachelor of Technology – Computer Science and Engineering; CGPA: 8.08"),
        ("Hindi Higher Secondary School — Tripura, India",
         "Intermediate – Percentage: 88% | Aug 2022"),
        ("Holy Cross School — Tripura, India",
         "Matriculation – Percentage: 94% | Aug 2020"),
    ]
    for name, detail in edu:
        p = para(doc, "", space_after=0)
        r = p.add_run(name); r.bold = True; r.font.size = Pt(10)
        para(doc, detail, size=9.5, italic=True, space_after=3)

    doc.save("cv/cv_template.docx")
    print("Wrote cv/cv_template.docx")


if __name__ == "__main__":
    main()
